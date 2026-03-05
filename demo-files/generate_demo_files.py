"""
generate_demo_files.py
Generates 5 realistic demo files with Indian PII for the hackathon demo.
Run from the demo-files/ directory, or any directory — output lands next to this script.
"""

import os
import random
import hashlib
from pathlib import Path

from faker import Faker

fake = Faker("en_IN")
Faker.seed(42)
random.seed(42)

OUT = Path(__file__).parent

# ── Helpers ───────────────────────────────────────────────────────────────────

def aadhaar() -> str:
    """Realistic 12-digit Aadhaar in XXXX XXXX XXXX format (first digit 2-9)."""
    first = random.randint(2, 9)
    rest = [random.randint(0, 9) for _ in range(11)]
    digits = [first] + rest
    return f"{''.join(map(str, digits[:4]))} {''.join(map(str, digits[4:8]))} {''.join(map(str, digits[8:]))}"

def pan() -> str:
    """Realistic PAN: 5 letters, 4 digits, 1 letter (e.g. ABCDE1234F)."""
    letters = "ABCDEFGHIJKLMNOPQRSTUVWXYZ"
    return (
        "".join(random.choices(letters, k=3))
        + random.choice("PCHABGJLFTE")   # 4th char = entity type
        + random.choice(letters)          # 5th = first letter of surname
        + "".join(str(random.randint(0, 9)) for _ in range(4))
        + random.choice(letters)
    )

def phone() -> str:
    """Indian mobile number starting with +91."""
    prefixes = ["7", "8", "9"]
    number = random.choice(prefixes) + "".join(str(random.randint(0, 9)) for _ in range(9))
    return f"+91 {number[:5]} {number[5:]}"

def card_number() -> str:
    """Fake Visa card number in XXXX XXXX XXXX XXXX format."""
    return f"4{''.join(str(random.randint(0,9)) for _ in range(3))} " \
           f"{''.join(str(random.randint(0,9)) for _ in range(4))} " \
           f"{''.join(str(random.randint(0,9)) for _ in range(4))} " \
           f"{''.join(str(random.randint(0,9)) for _ in range(4))}"

def cvv() -> str:
    return str(random.randint(100, 999))

def expiry() -> str:
    month = random.randint(1, 12)
    year = random.randint(26, 30)
    return f"{month:02d}/{year:02d}"

def upi_id(name: str) -> str:
    handle = name.lower().replace(" ", ".").split(".")[0]
    bank = random.choice(["okicici", "okhdfcbank", "okaxis", "paytm", "ybl"])
    return f"{handle}{random.randint(1,99)}@{bank}"

def ifsc() -> str:
    bank_code = random.choice(["HDFC", "ICIC", "SBIN", "AXIB", "KKBK"])
    return f"{bank_code}0{random.randint(100000, 999999)}"

def account_number() -> str:
    return "".join(str(random.randint(0, 9)) for _ in range(random.randint(10, 14)))

def passport() -> str:
    letters = "ABCDEFGHIJKLMNOPQRSTUVWXYZ"
    return random.choice(letters) + str(random.randint(1000000, 9999999))

def ip_address() -> str:
    return f"{random.randint(1,254)}.{random.randint(0,255)}.{random.randint(0,255)}.{random.randint(1,254)}"

def device_id() -> str:
    return ":".join(f"{random.randint(0,255):02X}" for _ in range(6))

def fingerprint_hash() -> str:
    return hashlib.sha256(fake.name().encode()).hexdigest()

# ── Person factory ────────────────────────────────────────────────────────────

CITIES = [
    "Mumbai", "Delhi", "Bengaluru", "Chennai", "Hyderabad",
    "Pune", "Kolkata", "Ahmedabad", "Jaipur", "Lucknow",
]

def make_person(idx: int) -> dict:
    first = fake.first_name()
    last = fake.last_name()
    full = f"{first} {last}"
    city = CITIES[idx % len(CITIES)]
    state_pin = {
        "Mumbai": ("Maharashtra", "400001"),
        "Delhi": ("Delhi", "110001"),
        "Bengaluru": ("Karnataka", "560001"),
        "Chennai": ("Tamil Nadu", "600001"),
        "Hyderabad": ("Telangana", "500001"),
        "Pune": ("Maharashtra", "411001"),
        "Kolkata": ("West Bengal", "700001"),
        "Ahmedabad": ("Gujarat", "380001"),
        "Jaipur": ("Rajasthan", "302001"),
        "Lucknow": ("Uttar Pradesh", "226001"),
    }
    state, pin = state_pin[city]
    address = f"{random.randint(1, 999)}, {fake.street_name()}, {city}, {state} - {pin}"
    return {
        "id": idx + 1,
        "first": first,
        "last": last,
        "full": full,
        "email": fake.email(),
        "phone": phone(),
        "aadhaar": aadhaar(),
        "pan": pan(),
        "card": card_number(),
        "cvv": cvv(),
        "expiry": expiry(),
        "address": address,
        "city": city,
        "upi": upi_id(full),
        "ifsc": ifsc(),
        "account": account_number(),
        "passport": passport(),
        "dob": fake.date_of_birth(minimum_age=22, maximum_age=55).strftime("%d/%m/%Y"),
        "bank": random.choice(["HDFC Bank", "ICICI Bank", "SBI", "Axis Bank", "Kotak Mahindra"]),
    }

people = [make_person(i) for i in range(5)]

# ── 1. demo_users.sql ─────────────────────────────────────────────────────────

def generate_sql():
    path = OUT / "demo_users.sql"
    lines = [
        "-- Demo users table with PII (for testing PII detection)",
        "CREATE TABLE IF NOT EXISTS users (",
        "    id           INT PRIMARY KEY,",
        "    first_name   VARCHAR(50),",
        "    last_name    VARCHAR(50),",
        "    email        VARCHAR(100),",
        "    phone        VARCHAR(20),",
        "    aadhaar      VARCHAR(14),",
        "    pan          VARCHAR(10),",
        "    card_number  VARCHAR(22),",
        "    cvv          VARCHAR(3),",
        "    expiry       VARCHAR(5)",
        ");",
        "",
        "INSERT INTO users (id, first_name, last_name, email, phone, aadhaar, pan, card_number, cvv, expiry) VALUES",
    ]
    rows = []
    for p in people:
        rows.append(
            f"  ({p['id']}, '{p['first']}', '{p['last']}', '{p['email']}', "
            f"'{p['phone']}', '{p['aadhaar']}', '{p['pan']}', "
            f"'{p['card']}', '{p['cvv']}', '{p['expiry']}')"
        )
    lines.append(",\n".join(rows) + ";")
    lines.append("")
    path.write_text("\n".join(lines), encoding="utf-8")
    pii_per_row = 8  # email, phone, aadhaar, pan, card, cvv, expiry + name
    return path.name, 5 * pii_per_row

# ── 2. demo_customers.csv ─────────────────────────────────────────────────────

def generate_csv():
    import csv
    path = OUT / "demo_customers.csv"
    headers = ["id", "full_name", "email", "phone", "address",
               "aadhaar", "pan", "upi_id", "ifsc", "account_number"]
    with open(path, "w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=headers)
        writer.writeheader()
        for p in people:
            writer.writerow({
                "id": p["id"],
                "full_name": p["full"],
                "email": p["email"],
                "phone": p["phone"],
                "address": p["address"],
                "aadhaar": p["aadhaar"],
                "pan": p["pan"],
                "upi_id": p["upi"],
                "ifsc": p["ifsc"],
                "account_number": p["account"],
            })
    pii_per_row = 9  # name, email, phone, address, aadhaar, pan, upi, ifsc, account
    return path.name, 5 * pii_per_row

# ── 3. demo_kyc.docx ──────────────────────────────────────────────────────────

def generate_docx():
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    p = people[0]  # single KYC form for person 0
    path = OUT / "demo_kyc.docx"
    doc = Document()

    # Title
    title = doc.add_heading("KYC Verification Form", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph("Know Your Customer — Individual")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()  # spacer

    # Section header
    def section(label: str):
        h = doc.add_paragraph()
        run = h.add_run(label)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        h.paragraph_format.space_before = Pt(8)

    # KYC table
    section("Personal Information")
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.columns[0].width = Inches(2.2)
    table.columns[1].width = Inches(4.0)

    def add_row(label: str, value: str):
        row = table.add_row()
        lbl_cell = row.cells[0]
        val_cell = row.cells[1]
        lbl_run = lbl_cell.paragraphs[0].add_run(label)
        lbl_run.bold = True
        lbl_run.font.size = Pt(10)
        val_run = val_cell.paragraphs[0].add_run(value)
        val_run.font.size = Pt(10)

    add_row("Full Name:", p["full"])
    add_row("Date of Birth:", p["dob"])
    add_row("Aadhaar Number:", p["aadhaar"])
    add_row("PAN:", p["pan"])
    add_row("Passport Number:", p["passport"])
    add_row("Mobile:", p["phone"])
    add_row("Email:", p["email"])
    add_row("Residential Address:", p["address"])

    doc.add_paragraph()
    section("Bank Details")

    table2 = doc.add_table(rows=0, cols=2)
    table2.style = "Table Grid"
    table2.columns[0].width = Inches(2.2)
    table2.columns[1].width = Inches(4.0)

    def add_row2(label: str, value: str):
        row = table2.add_row()
        lbl_run = row.cells[0].paragraphs[0].add_run(label)
        lbl_run.bold = True
        lbl_run.font.size = Pt(10)
        row.cells[1].paragraphs[0].add_run(value).font.size = Pt(10)

    add_row2("Bank Name:", p["bank"])
    add_row2("Account Number:", p["account"])
    add_row2("IFSC Code:", p["ifsc"])
    add_row2("UPI ID:", p["upi"])

    doc.add_paragraph()
    doc.add_paragraph(
        "I hereby declare that the information furnished above is true, correct and complete "
        "to the best of my knowledge and belief."
    )
    sig = doc.add_paragraph()
    sig.add_run(f"\nSignature: ________________      Date: {p['dob']}").italic = True

    doc.save(path)
    # PII: full name, dob, aadhaar, pan, passport, phone, email, address, account, ifsc, upi, bank = 12
    return path.name, 12

# ── 4. demo_notes.txt ─────────────────────────────────────────────────────────

def generate_txt():
    p1, p2, p3 = people[0], people[1], people[2]
    path = OUT / "demo_notes.txt"

    text = f"""Meeting Notes — Product & Compliance Sync
Date: {fake.date_this_year().strftime("%d %B %Y")}
Attendees: {p1['full']} (Product), {p2['full']} (Compliance)

Discussion Summary:

{p1['full']} raised a concern about the onboarding flow for new customers in Tier-2 cities. \
The latest batch of signups included a user who registered with mobile number {p1['phone']} \
and email {p1['email']}. During KYC verification, the submitted Aadhaar {p1['aadhaar']} \
was flagged by the UIDAI API with a soft mismatch — the address on record listed \
"{p1['address']}" but the submitted address differed in the pincode field. \
The discrepancy was minor and the case was escalated to the compliance team.

{p2['full']} confirmed that the compliance portal logged the event under reference PAN {p1['pan']}. \
She mentioned that an automated alert was sent to the registered UPI handle {p1['upi']}. \
The system also captured the request originating from IP address {ip_address()}, \
device MAC {device_id()}, and the session was fingerprinted as \
{fingerprint_hash()[:40]}. {p2['first']} also noted that all such events are \
stored with a 90-day retention policy as per RBI guidelines.

Action items for next sprint: integrate the UIDAI sandbox for soft-match resolution, \
update the KYC microservice to store device fingerprints separately from PII fields, \
and schedule a review with {p3['full']} from legal to confirm data retention timelines. \
{p3['first']} can be reached at {p3['phone']} for a quick pre-sprint call.
"""
    path.write_text(text.strip(), encoding="utf-8")
    # PII: 3 names, 2 phones, 1 email, 1 aadhaar, 1 pan, 1 address, 1 ip, 1 upi, 1 device_id, 1 fingerprint = 13
    return path.name, 13

# ── 5. demo_report.pdf ────────────────────────────────────────────────────────

def generate_pdf():
    from fpdf import FPDF

    p = people[3]
    path = OUT / "demo_report.pdf"

    pdf = FPDF()
    pdf.add_page()

    # Header band
    pdf.set_fill_color(31, 73, 125)
    pdf.rect(0, 0, 210, 30, style="F")
    pdf.set_text_color(255, 255, 255)
    pdf.set_font("Helvetica", "B", 16)
    pdf.set_xy(10, 8)
    pdf.cell(0, 12, "HR Employee Record - Confidential", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 9)
    pdf.set_xy(10, 20)
    pdf.cell(0, 8, f"Generated: {fake.date_this_year().strftime('%d %B %Y')}  |  Department: Human Resources")

    # Reset color
    pdf.set_text_color(0, 0, 0)
    pdf.set_xy(10, 38)

    pdf.set_font("Helvetica", "B", 13)
    pdf.cell(0, 10, "Employee Details", new_x="LMARGIN", new_y="NEXT")

    # Divider
    pdf.set_draw_color(200, 200, 200)
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(4)

    def field(label: str, value: str):
        pdf.set_font("Helvetica", "B", 10)
        pdf.set_x(10)
        pdf.cell(55, 8, label, border=0)
        pdf.set_font("Helvetica", "", 10)
        pdf.cell(0, 8, value, new_x="LMARGIN", new_y="NEXT")

    field("Full Name:", p["full"])
    field("Employee ID:", f"EMP{random.randint(10000,99999)}")
    field("Email Address:", p["email"])
    field("Mobile:", p["phone"])
    field("Date of Birth:", p["dob"])
    field("Aadhaar Number:", p["aadhaar"])
    field("PAN:", p["pan"])
    field("Residential Address:", p["address"])

    pdf.ln(3)
    pdf.set_draw_color(200, 200, 200)
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(4)

    pdf.set_font("Helvetica", "B", 13)
    pdf.cell(0, 10, "Bank & Payment Details", new_x="LMARGIN", new_y="NEXT")
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(4)

    field("Bank Name:", p["bank"])
    field("Account Number:", p["account"])
    field("IFSC Code:", p["ifsc"])
    field("UPI ID:", p["upi"])

    pdf.ln(6)
    pdf.set_font("Helvetica", "I", 8)
    pdf.set_text_color(120, 120, 120)
    pdf.multi_cell(
        0, 5,
        "This document contains sensitive personal and financial information. "
        "Unauthorized access, reproduction or distribution is strictly prohibited under the IT Act 2000 and DPDPA 2023."
    )

    pdf.output(str(path))
    # PII: name, email, phone, dob, aadhaar, pan, address, account, ifsc, upi, bank = 11
    return path.name, 11

# ── Main ──────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    results = [
        generate_sql(),
        generate_csv(),
        generate_docx(),
        generate_txt(),
        generate_pdf(),
    ]

    print("\nGenerated files:")
    total = 0
    for name, count in results:
        print(f"  {name:<25}  {count} PII instances embedded")
        total += count
    print(f"\n  Total PII across all files: {total}")
    print(f"\n  Output directory: {OUT.resolve()}")
