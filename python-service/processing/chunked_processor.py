"""
Chunked parallel processing for large files (> LARGE_FILE_THRESHOLD bytes).

Each file format is split at logical boundaries (never byte boundaries), chunks
are processed in parallel with ThreadPoolExecutor, and the results are merged
back into the original file format.

Context overlap (CONTEXT_OVERLAP chars) is added to either side of each text
chunk so that PII spanning a chunk boundary is still detected.  The overlap
is used for detection only — replacement is applied to the chunk's own content.

Chunk targets
─────────────
  SQL   : 500 statements per chunk
  CSV   : 10 000 rows per chunk
  TXT   : 50 000 characters per chunk (split at paragraph boundaries)
  JSON  : 1 000 items per chunk (top-level array only)
  PDF   : 10 pages per chunk
  DOCX  : 200 paragraphs per chunk
  Image : 4×4 grid = 16 tiles

Parallelism
───────────
  ThreadPoolExecutor with min(4, len(chunks)) workers.
  PDF and DOCX: detection parallel, document mutation single-threaded
                (PyMuPDF and python-docx are not thread-safe for writes).
"""

from __future__ import annotations

import json
import re
import threading
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Callable, Optional

from detection.analyzer_engine import pii_analyzer
from detection.context_analyzer import context_analyzer
from detection.confidence_scorer import confidence_scorer
from detection.masker import pii_masker

# ── Thresholds and targets ─────────────────────────────────────────────────────

LARGE_FILE_THRESHOLD     = 5 * 1024 * 1024   # 5 MB — use chunked above this
SQL_STATEMENTS_PER_CHUNK = 500
CSV_ROWS_PER_CHUNK       = 2_000    # each row-chunk processed in ~10-20 s on CPU
TXT_CHARS_PER_CHUNK      = 30_000   # ~5-8 s per chunk
JSON_ITEMS_PER_CHUNK     = 300
PDF_PAGES_PER_CHUNK      = 5
DOCX_PARAS_PER_CHUNK     = 100
GRID_SIZE                = 4                  # 4×4 = 16 image tiles
CONTEXT_OVERLAP          = 500               # chars shared with neighbours
_CSV_CELL_BATCH          = 100               # rows per NLP call — higher is fine without BERT (spaCy scales linearly)
_MAX_WORKERS             = 6                 # ThreadPoolExecutor threads; BERT releases GIL

ProgressCallback = Callable[[int, str], None]   # (chunk_idx, status: "pending"|"processing"|"done"|"failed")


def _noop_progress(chunk_idx: int, status: str) -> None:  # noqa: ARG001
    pass


# ── Internal data types ────────────────────────────────────────────────────────

@dataclass
class _ChunkResult:
    index:       int
    masked_text: str                  = ""
    to_mask:     list[dict[str, Any]] = field(default_factory=list)
    high_count:  int                  = 0
    medium_count: int                 = 0


# ── Core pipeline helpers ──────────────────────────────────────────────────────

def _pipeline_value_map(
    text: str,
    mode: str,
    column_name: str | None = None,
    skip_transformer: bool = False,
) -> tuple[dict[str, str], list[dict[str, Any]], int, int]:
    """
    Run the 5-stage detection pipeline on *text* and return:
      - replacement_map : {pii_value: masked_replacement}
      - to_mask         : raw detection results
      - high_count, medium_count

    Value-based replacement (rather than position-based) lets the same map
    be applied to any subset of the original text.

    skip_transformer=True skips the BERT layer — use for CSV/JSON where
    structured PII is fully covered by regex + spaCy but BERT costs ~2 s/call.
    """
    analysis = pii_analyzer.analyze(text, skip_transformer=skip_transformer)
    enriched = context_analyzer.analyze(
        analysis["cleaned_text"],
        analysis["presidio_results"],
        analysis["indic_results"],
        analysis["label_pairs"],
        column_name=column_name,
    )
    deduped = confidence_scorer.deduplicate(enriched)
    scored  = confidence_scorer.score_and_filter(deduped)
    to_mask = scored["to_mask"]

    replacement_map: dict[str, str] = {}
    for result in to_mask:
        value = result.get("value", "")
        if not value or value in replacement_map:
            continue
        single = pii_masker.mask(value, [{**result, "start": 0, "end": len(value)}], mode)
        replacement_map[value] = single["masked_text"]

    return replacement_map, to_mask, scored["high_count"], scored["medium_count"]


def _apply_map(text: str, replacement_map: dict[str, str]) -> str:
    """Apply a value replacement map, longest keys first to avoid partial matches."""
    for original in sorted(replacement_map, key=len, reverse=True):
        text = text.replace(original, replacement_map[original])
    return text


def _run_with_overlap(
    before:  str,
    content: str,
    after:   str,
    mode:    str,
) -> tuple[str, list[dict[str, Any]], int, int]:
    """
    Detect PII in *before + content + after* (wider context window) then
    apply the replacement map to *content* only.  This ensures PII at
    chunk boundaries is caught without duplicating output.
    """
    window = before + content + after
    replacement_map, to_mask, high, medium = _pipeline_value_map(window, mode)
    return _apply_map(content, replacement_map), to_mask, high, medium


def _merge_stats(
    results: list[_ChunkResult],
) -> tuple[list[dict[str, Any]], int, int]:
    all_to_mask:  list[dict[str, Any]] = []
    total_high   = 0
    total_medium = 0
    for r in results:
        all_to_mask.extend(r.to_mask)
        total_high   += r.high_count
        total_medium += r.medium_count
    return all_to_mask, total_high, total_medium


def _build_summary(
    all_to_mask:  list[dict[str, Any]],
    high:         int,
    medium:       int,
    chunk_count:  int,
) -> dict[str, Any]:
    return {
        "pii_summary":          confidence_scorer.get_summary(all_to_mask),
        "layer_breakdown":      confidence_scorer.get_layer_breakdown(all_to_mask),
        "confidence_breakdown": {"high": high, "medium": medium},
        "strategies_applied":   {},
        "total_pii":            len(all_to_mask),
        "chunk_count":          chunk_count,
    }


def _parallel_text_run(
    raw_chunks:  list[str],
    mode:        str,
    progress_cb: ProgressCallback,
    use_overlap: bool = True,
) -> list[_ChunkResult]:
    """
    Process text chunks concurrently via ThreadPoolExecutor.

    PyTorch / spaCy C-extensions release the GIL during inference, so multiple
    threads genuinely overlap during the BERT forward pass.  Results are
    collected in index order so the output is always deterministic.
    """
    total   = len(raw_chunks)
    results: dict[int, _ChunkResult] = {}

    for i in range(total):
        progress_cb(i, "pending")

    def worker(idx: int, content: str) -> _ChunkResult:
        progress_cb(idx, "processing")
        try:
            if use_overlap and total > 1:
                before = raw_chunks[idx - 1][-CONTEXT_OVERLAP:] if idx > 0         else ""
                after  = raw_chunks[idx + 1][:CONTEXT_OVERLAP]  if idx < total - 1 else ""
                masked, to_mask, high, medium = _run_with_overlap(before, content, after, mode)
            else:
                repl_map, to_mask, high, medium = _pipeline_value_map(content, mode)
                masked = _apply_map(content, repl_map)
            result = _ChunkResult(idx, masked, to_mask, high, medium)
            progress_cb(idx, "done")
            return result
        except Exception:
            progress_cb(idx, "failed")
            raise

    with ThreadPoolExecutor(max_workers=min(_MAX_WORKERS, total)) as executor:
        futures = {executor.submit(worker, i, chunk): i for i, chunk in enumerate(raw_chunks)}
        for future in as_completed(futures):
            r = future.result()
            results[r.index] = r

    return [results[i] for i in range(total)]


# ── SQL ───────────────────────────────────────────────────────────────────────

def _split_sql_statements(text: str) -> list[str]:
    """Split SQL into individual statements at semicolons, preserving whitespace."""
    parts      = re.split(r"(;)", text)
    statements: list[str] = []
    current    = ""
    for part in parts:
        current += part
        if part == ";" and current.strip() not in ("", ";"):
            statements.append(current)
            current = ""
    if current.strip():
        statements.append(current)
    return statements or [text]


def process_sql_chunked(
    input_path:  str,
    output_path: str,
    mode:        str              = "redact",
    progress_cb: ProgressCallback = _noop_progress,
) -> dict[str, Any]:
    with open(input_path, encoding="utf-8", errors="replace") as fh:
        text = fh.read()

    statements  = _split_sql_statements(text)
    raw_chunks  = [
        "".join(statements[i : i + SQL_STATEMENTS_PER_CHUNK])
        for i in range(0, len(statements), SQL_STATEMENTS_PER_CHUNK)
    ] or [text]

    ordered = _parallel_text_run(raw_chunks, mode, progress_cb)

    with open(output_path, "w", encoding="utf-8") as fh:
        fh.write("".join(r.masked_text for r in ordered))

    all_to_mask, high, medium = _merge_stats(ordered)
    return _build_summary(all_to_mask, high, medium, len(raw_chunks))


# ── CSV ───────────────────────────────────────────────────────────────────────

def process_csv_chunked(
    input_path:  str,
    output_path: str,
    mode:        str              = "redact",
    progress_cb: ProgressCallback = _noop_progress,
) -> dict[str, Any]:
    import pandas as pd

    df          = pd.read_csv(input_path, dtype=str, keep_default_na=False)
    row_groups  = [
        df.iloc[i : i + CSV_ROWS_PER_CHUNK].copy()
        for i in range(0, max(1, len(df)), CSV_ROWS_PER_CHUNK)
    ]
    total        = len(row_groups)
    all_to_mask: list[dict[str, Any]] = []
    all_to_mask_lock = threading.Lock()
    total_high   = 0
    total_medium = 0
    write_lock   = threading.Lock()
    write_header_flag = [True]   # mutable container so closure can update it

    for i in range(total):
        progress_cb(i, "pending")

    def process_chunk(idx: int, chunk_df):
        """Process one row-group: columns are analysed in parallel, then the
        sanitised DataFrame is stream-appended to the output CSV."""
        progress_cb(idx, "processing")
        try:
            result_df    = chunk_df.copy()
            col_results: dict[str, list[str]] = {}
            col_to_mask: list[dict[str, Any]] = []
            col_lock     = threading.Lock()

            def process_col(col: str):
                values        = chunk_df[col].tolist()
                col_sanitised: list[str] = []
                local_mask:   list[dict[str, Any]] = []
                for batch_start in range(0, len(values), _CSV_CELL_BATCH):
                    batch   = values[batch_start: batch_start + _CSV_CELL_BATCH]
                    joined  = " | ".join(str(v) for v in batch)
                    # skip_transformer=True: BERT adds ~2s/call on CPU but CSV
                    # column values are structured — regex + spaCy catch all of
                    # emails, phones, SSNs, Aadhaar, PAN, credit cards, IPs.
                    repl_map, to_mask, _, _ = _pipeline_value_map(
                        joined, mode, column_name=col, skip_transformer=True
                    )
                    local_mask.extend(to_mask)
                    col_sanitised.extend(_apply_map(str(v), repl_map) for v in batch)
                with col_lock:
                    col_results[col] = col_sanitised
                    col_to_mask.extend(local_mask)

            # Parallelise across columns — each column's NLP calls are independent
            cols = list(chunk_df.columns)
            with ThreadPoolExecutor(max_workers=min(_MAX_WORKERS, len(cols))) as col_exec:
                col_futs = [col_exec.submit(process_col, c) for c in cols]
                for f in as_completed(col_futs):
                    f.result()  # re-raise any exception

            for col in cols:
                result_df[col] = col_results[col]

            # Stream-append results in chunk order under a lock so CSV rows stay ordered
            with write_lock:
                is_first = write_header_flag[0]
                result_df.to_csv(
                    output_path,
                    mode="w" if is_first else "a",
                    header=is_first,
                    index=False,
                )
                write_header_flag[0] = False

            high   = sum(1 for r in col_to_mask if r.get("confidence") == "high")
            medium = sum(1 for r in col_to_mask if r.get("confidence") == "medium")
            progress_cb(idx, "done")
            return col_to_mask, high, medium
        except Exception:
            progress_cb(idx, "failed")
            raise

    # Process row-chunks sequentially to preserve CSV row order in the output
    # (columns within each chunk are still parallelised above)
    for idx, grp in enumerate(row_groups):
        to_mask, high, medium = process_chunk(idx, grp)
        with all_to_mask_lock:
            all_to_mask.extend(to_mask)
            total_high   += high
            total_medium += medium

    return _build_summary(all_to_mask, total_high, total_medium, total)


# ── TXT ───────────────────────────────────────────────────────────────────────

def _split_txt_paragraphs(text: str) -> list[str]:
    """Split text at paragraph boundaries (double newlines) into ~TXT_CHARS_PER_CHUNK groups."""
    parts   = re.split(r"(\n\n+)", text)
    chunks: list[str] = []
    current = ""
    for part in parts:
        if len(current) + len(part) > TXT_CHARS_PER_CHUNK and current:
            chunks.append(current)
            current = part
        else:
            current += part
    if current:
        chunks.append(current)
    return chunks or [text]


def process_txt_chunked(
    input_path:  str,
    output_path: str,
    mode:        str              = "redact",
    progress_cb: ProgressCallback = _noop_progress,
) -> dict[str, Any]:
    with open(input_path, encoding="utf-8", errors="replace") as fh:
        text = fh.read()

    raw_chunks = _split_txt_paragraphs(text)
    ordered    = _parallel_text_run(raw_chunks, mode, progress_cb)

    with open(output_path, "w", encoding="utf-8") as fh:
        fh.write("".join(r.masked_text for r in ordered))

    all_to_mask, high, medium = _merge_stats(ordered)
    return _build_summary(all_to_mask, high, medium, len(raw_chunks))


# ── JSON ──────────────────────────────────────────────────────────────────────

def _sanitise_json_node(
    node:        Any,
    mode:        str,
    all_to_mask: list[dict[str, Any]],
    counters:    dict[str, int],
) -> Any:
    """Recursively sanitise all string leaves (used for non-array JSON roots)."""
    if isinstance(node, dict):
        return {k: _sanitise_json_node(v, mode, all_to_mask, counters) for k, v in node.items()}
    if isinstance(node, list):
        return [_sanitise_json_node(item, mode, all_to_mask, counters) for item in node]
    if isinstance(node, str) and node.strip():
        analysis = pii_analyzer.analyze(node)
        enriched = context_analyzer.analyze(
            analysis["cleaned_text"],
            analysis["presidio_results"],
            analysis["indic_results"],
            analysis["label_pairs"],
        )
        deduped = confidence_scorer.deduplicate(enriched)
        scored  = confidence_scorer.score_and_filter(deduped)
        to_mask = scored["to_mask"]
        all_to_mask.extend(to_mask)
        counters["high"]   = counters.get("high",   0) + scored["high_count"]
        counters["medium"] = counters.get("medium", 0) + scored["medium_count"]
        if not to_mask:
            return node
        return pii_masker.mask(analysis["cleaned_text"], to_mask, mode)["masked_text"]
    return node


def process_json_chunked(
    input_path:  str,
    output_path: str,
    mode:        str              = "redact",
    progress_cb: ProgressCallback = _noop_progress,
) -> dict[str, Any]:
    with open(input_path, encoding="utf-8") as fh:
        data = json.load(fh)

    with open(input_path, encoding="utf-8") as fh:
        raw_head = fh.read(512)
    indent = 2 if ("\n  " in raw_head or "\n\t" in raw_head) else None

    if isinstance(data, list):
        item_groups = [
            data[i : i + JSON_ITEMS_PER_CHUNK]
            for i in range(0, max(1, len(data)), JSON_ITEMS_PER_CHUNK)
        ]
        total          = len(item_groups)
        merged_items: list[Any]       = []
        chunk_results: dict[int, tuple[list, list[dict], int, int]] = {}
        results_lock   = threading.Lock()

        for i in range(total):
            progress_cb(i, "pending")

        def process_group(idx: int, items: list):
            progress_cb(idx, "processing")
            try:
                local_to_mask: list[dict[str, Any]] = []
                counters:      dict[str, int]        = {}
                sanitised = [_sanitise_json_node(item, mode, local_to_mask, counters) for item in items]
                progress_cb(idx, "done")
                return idx, sanitised, local_to_mask, counters.get("high", 0), counters.get("medium", 0)
            except Exception:
                progress_cb(idx, "failed")
                raise

        with ThreadPoolExecutor(max_workers=min(_MAX_WORKERS, total)) as executor:
            futures = {executor.submit(process_group, i, grp): i for i, grp in enumerate(item_groups)}
            for future in as_completed(futures):
                idx, sanitised, to_mask, high, medium = future.result()
                with results_lock:
                    chunk_results[idx] = (sanitised, to_mask, high, medium)

        all_to_mask:  list[dict] = []
        total_high   = 0
        total_medium = 0
        for i in range(total):
            sanitised, to_mask, high, medium = chunk_results[i]
            merged_items.extend(sanitised)
            all_to_mask.extend(to_mask)
            total_high   += high
            total_medium += medium
        output_data = merged_items

    else:
        all_to_mask:  list[dict[str, Any]] = []
        counters:     dict[str, int]        = {}
        progress_cb(0, "pending")
        progress_cb(0, "processing")
        output_data  = _sanitise_json_node(data, mode, all_to_mask, counters)
        total_high   = counters.get("high",   0)
        total_medium = counters.get("medium", 0)
        total        = 1
        progress_cb(0, "done")

    with open(output_path, "w", encoding="utf-8") as fh:
        json.dump(output_data, fh, indent=indent, ensure_ascii=False)

    return _build_summary(all_to_mask, total_high, total_medium, total)


# ── PDF ───────────────────────────────────────────────────────────────────────

def process_pdf_chunked(
    input_path:  str,
    output_path: str,
    mode:        str              = "redact",
    progress_cb: ProgressCallback = _noop_progress,
) -> dict[str, Any]:
    import fitz  # PyMuPDF

    doc        = fitz.open(input_path)
    page_count = len(doc)

    # Collect page texts in the main thread — PyMuPDF is not thread-safe
    all_page_texts: list[str] = []
    for page in doc:
        text = page.get_text()
        try:
            import io
            import pytesseract
            from PIL import Image as PILImage
            for img_info in page.get_images(full=True):
                xref    = img_info[0]
                base_img = doc.extract_image(xref)
                img     = PILImage.open(io.BytesIO(base_img["image"]))
                ocr     = pytesseract.image_to_string(img)
                if ocr.strip():
                    text += "\n" + ocr
        except Exception:
            pass
        all_page_texts.append(text)

    page_groups: list[list[int]] = [
        list(range(i, min(i + PDF_PAGES_PER_CHUNK, page_count)))
        for i in range(0, max(1, page_count), PDF_PAGES_PER_CHUNK)
    ]
    total          = len(page_groups)
    lock           = threading.Lock()
    all_pii_values: set[str]          = set()
    all_to_mask:   list[dict[str, Any]] = []
    total_high     = 0
    total_medium   = 0

    for i in range(total):
        progress_cb(i, "pending")

    def detect_group(idx: int, page_indices: list[int]):
        progress_cb(idx, "processing")
        try:
            group_text = "\n".join(all_page_texts[i] for i in page_indices)
            analysis   = pii_analyzer.analyze(group_text)
            enriched   = context_analyzer.analyze(
                analysis["cleaned_text"],
                analysis["presidio_results"],
                analysis["indic_results"],
                analysis["label_pairs"],
            )
            deduped = confidence_scorer.deduplicate(enriched)
            scored  = confidence_scorer.score_and_filter(deduped)
            progress_cb(idx, "done")
            return scored["to_mask"], scored["high_count"], scored["medium_count"]
        except Exception:
            progress_cb(idx, "failed")
            raise

    with ThreadPoolExecutor(max_workers=min(_MAX_WORKERS, total)) as executor:
        futures = {executor.submit(detect_group, i, grp): i for i, grp in enumerate(page_groups)}
        for future in as_completed(futures):
            try:
                to_mask, high, medium = future.result()
                with lock:
                    all_to_mask.extend(to_mask)
                    total_high   += high
                    total_medium += medium
                    for r in to_mask:
                        v = r.get("value", "")
                        if v:
                            all_pii_values.add(v)
            except Exception:
                pass  # progress already set to failed inside detect_group

    # Apply redactions single-threaded (PyMuPDF write requirement)
    for page in doc:
        for pii_value in all_pii_values:
            for area in page.search_for(pii_value):
                page.add_redact_annot(area, fill=(0, 0, 0))
        page.apply_redactions()

    for field_name in ("author", "creator", "producer", "subject", "title"):
        try:
            doc.set_metadata({field_name: ""})
        except Exception:
            pass

    doc.save(output_path, garbage=4, deflate=True)
    doc.close()

    return _build_summary(all_to_mask, total_high, total_medium, total)


# ── DOCX ──────────────────────────────────────────────────────────────────────

def _collect_all_paragraphs(doc: Any) -> list[Any]:
    """Return every paragraph from body, tables, headers, and footers."""
    paras: list[Any] = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paras.extend(cell.paragraphs)
    for section in doc.sections:
        for container in (section.header, section.footer):
            if not container.is_linked_to_previous:
                paras.extend(container.paragraphs)
    return paras


def _replace_para(para: Any, replacement_map: dict[str, str]) -> None:
    """Apply replacement_map to a paragraph, handling cross-run PII."""
    if not para.runs:
        return
    # Per-run pass (preserves formatting for within-run PII)
    for run in para.runs:
        for original, replacement in replacement_map.items():
            if original in run.text:
                run.text = run.text.replace(original, replacement)
    # Cross-run pass (merges runs for PII that spans run boundaries)
    full     = "".join(r.text for r in para.runs)
    modified = _apply_map(full, replacement_map)
    if modified != full and para.runs:
        para.runs[0].text = modified
        for run in para.runs[1:]:
            run.text = ""


def process_docx_chunked(
    input_path:  str,
    output_path: str,
    mode:        str              = "redact",
    progress_cb: ProgressCallback = _noop_progress,
) -> dict[str, Any]:
    from docx import Document

    doc   = Document(input_path)
    paras = _collect_all_paragraphs(doc)

    para_groups: list[list[Any]] = [
        paras[i : i + DOCX_PARAS_PER_CHUNK]
        for i in range(0, max(1, len(paras)), DOCX_PARAS_PER_CHUNK)
    ]
    total        = len(para_groups)
    lock         = threading.Lock()
    det_results: dict[int, tuple[dict[str, str], list[dict], int, int]] = {}
    all_to_mask: list[dict[str, Any]] = []
    total_high   = 0
    total_medium = 0

    for i in range(total):
        progress_cb(i, "pending")

    def detect_group(idx: int, group: list[Any]):
        progress_cb(idx, "processing")
        try:
            chunk_text = "\n".join(" ".join(r.text for r in p.runs) for p in group)
            if not chunk_text.strip():
                progress_cb(idx, "done")
                return idx, {}, [], 0, 0
            repl_map, to_mask, high, medium = _pipeline_value_map(chunk_text, mode)
            progress_cb(idx, "done")
            return idx, repl_map, to_mask, high, medium
        except Exception:
            progress_cb(idx, "failed")
            raise

    # Detection in parallel (python-docx reads are thread-safe)
    with ThreadPoolExecutor(max_workers=min(_MAX_WORKERS, total)) as executor:
        futures = {executor.submit(detect_group, i, grp): i for i, grp in enumerate(para_groups)}
        for future in as_completed(futures):
            chunk_idx = futures[future]
            try:
                idx, repl_map, to_mask, high, medium = future.result()
                with lock:
                    det_results[idx] = (repl_map, to_mask, high, medium)
            except Exception:
                with lock:
                    det_results[chunk_idx] = ({}, [], 0, 0)

    # Apply replacements single-threaded (python-docx writes are NOT thread-safe)
    for i, group in enumerate(para_groups):
        repl_map, to_mask, high, medium = det_results[i]
        all_to_mask.extend(to_mask)
        total_high   += high
        total_medium += medium
        if repl_map:
            for para in group:
                _replace_para(para, repl_map)

    core = doc.core_properties
    for attr in ("author", "last_modified_by", "comments"):
        try:
            setattr(core, attr, "Sanitized")
        except Exception:
            pass

    doc.save(output_path)
    return _build_summary(all_to_mask, total_high, total_medium, total)


# ── Image (4×4 grid tiles) ────────────────────────────────────────────────────

def process_image_chunked(
    input_path:  str,
    output_path: str,
    mode:        str              = "redact",   # noqa: ARG001 — redaction only for images
    progress_cb: ProgressCallback = _noop_progress,
) -> dict[str, Any]:
    try:
        import pytesseract
        from PIL import Image, ImageDraw
    except ImportError as exc:
        raise RuntimeError("pytesseract and Pillow are required for image processing.") from exc

    img          = Image.open(input_path).convert("RGB")
    width, height = img.size
    tile_w = width  // GRID_SIZE
    tile_h = height // GRID_SIZE

    # (tile_index, x0, y0, x1, y1)
    tiles: list[tuple[int, int, int, int, int]] = [
        (
            row * GRID_SIZE + col,
            col * tile_w,
            row * tile_h,
            (col + 1) * tile_w if col < GRID_SIZE - 1 else width,
            (row + 1) * tile_h if row < GRID_SIZE - 1 else height,
        )
        for row in range(GRID_SIZE)
        for col in range(GRID_SIZE)
    ]
    total          = len(tiles)
    lock           = threading.Lock()
    all_pii_values: set[str]           = set()
    all_to_mask:   list[dict[str, Any]] = []
    total_high     = 0
    total_medium   = 0

    for i in range(total):
        progress_cb(i, "pending")

    def detect_tile(tile_info: tuple[int, int, int, int, int]):
        tile_idx, x0, y0, x1, y1 = tile_info
        progress_cb(tile_idx, "processing")
        try:
            tile_img  = img.crop((x0, y0, x1, y1))
            tile_text = pytesseract.image_to_string(tile_img)
            if not tile_text.strip():
                progress_cb(tile_idx, "done")
                return [], 0, 0
            analysis = pii_analyzer.analyze(tile_text)
            enriched = context_analyzer.analyze(
                analysis["cleaned_text"],
                analysis["presidio_results"],
                analysis["indic_results"],
                analysis["label_pairs"],
            )
            deduped = confidence_scorer.deduplicate(enriched)
            scored  = confidence_scorer.score_and_filter(deduped)
            progress_cb(tile_idx, "done")
            return scored["to_mask"], scored["high_count"], scored["medium_count"]
        except Exception:
            progress_cb(tile_idx, "failed")
            raise

    with ThreadPoolExecutor(max_workers=min(_MAX_WORKERS, total)) as executor:
        futures = {executor.submit(detect_tile, t): t[0] for t in tiles}
        for future in as_completed(futures):
            try:
                to_mask, high, medium = future.result()
                with lock:
                    all_to_mask.extend(to_mask)
                    total_high   += high
                    total_medium += medium
                    for r in to_mask:
                        v = r.get("value", "")
                        if v:
                            all_pii_values.add(v)
            except Exception:
                pass  # progress set to failed inside detect_tile

    # Redact on the full image using word-level OCR bounding boxes
    ocr_data = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT)
    draw     = ImageDraw.Draw(img)
    n        = len(ocr_data["text"])
    for i in range(n):
        word = (ocr_data["text"][i] or "").strip()
        if not word:
            continue
        for pii_value in all_pii_values:
            if word.lower() in pii_value.lower() or pii_value.lower() in word.lower():
                x, y, w, h = (
                    ocr_data["left"][i],
                    ocr_data["top"][i],
                    ocr_data["width"][i],
                    ocr_data["height"][i],
                )
                draw.rectangle([x, y, x + w, y + h], fill=(0, 0, 0))
                break

    img.save(output_path)
    return _build_summary(all_to_mask, total_high, total_medium, total)


# ── Public dispatch table ─────────────────────────────────────────────────────

CHUNKED_FUNCS: dict[str, Any] = {
    "sql":  process_sql_chunked,
    "csv":  process_csv_chunked,
    "txt":  process_txt_chunked,
    "md":   process_txt_chunked,
    "json": process_json_chunked,
    "pdf":  process_pdf_chunked,
    "docx": process_docx_chunked,
    "doc":  process_docx_chunked,
    "png":  process_image_chunked,
    "jpg":  process_image_chunked,
    "jpeg": process_image_chunked,
}
