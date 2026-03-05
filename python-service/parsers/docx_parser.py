"""
DOCX parser using python-docx.

Collects text from paragraphs, table cells, headers, and footers, runs the
full detection pipeline, then applies in-place string replacements on every
Run object (preserving formatting).  Document metadata (author, last_modified_by)
is cleared.
"""

from __future__ import annotations

from typing import Any

from docx import Document  # type: ignore[import]
from docx.oxml.ns import qn  # type: ignore[import]

from detection.analyzer_engine import pii_analyzer
from detection.context_analyzer import context_analyzer
from detection.confidence_scorer import confidence_scorer
from detection.masker import pii_masker


def _collect_runs(doc: Document) -> list[Any]:
    """Return every Run object from paragraphs, tables, headers, and footers."""
    runs: list[Any] = []

    for para in doc.paragraphs:
        runs.extend(para.runs)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    runs.extend(para.runs)

    for section in doc.sections:
        for container in (section.header, section.footer):
            if container.is_linked_to_previous:
                continue
            for para in container.paragraphs:
                runs.extend(para.runs)

    return runs


def _collect_paragraphs(doc: Document) -> list[Any]:
    """Return every paragraph from the document including tables, headers, footers."""
    paras: list[Any] = []

    paras.extend(doc.paragraphs)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paras.extend(cell.paragraphs)

    for section in doc.sections:
        for container in (section.header, section.footer):
            if container.is_linked_to_previous:
                continue
            paras.extend(container.paragraphs)

    return paras


def _replace_in_paragraph(para: Any, replacement_map: dict[str, str]) -> None:
    """
    Replace PII in a paragraph, handling values that span multiple runs.

    Strategy:
    1. Per-run pass: replace PII values that appear entirely within a single run
       (preserves per-run formatting for non-spanning PII).
    2. Cross-run pass: join all run texts, replace any remaining PII values,
       and if the merged text changed, put it into the first run and blank the
       others (keeps the first run's character formatting).
    3. Verification: re-read each run after step 1 to confirm the replacement
       actually took effect; if not, the cross-run pass in step 2 covers it.
    """
    if not para.runs:
        return

    # ── Step 1: per-run replacement ───────────────────────────────────────────
    for run in para.runs:
        original_text = run.text
        for original, replacement in replacement_map.items():
            if original in run.text:
                run.text = run.text.replace(original, replacement)
        # Verify replacement was applied (re-read to confirm)
        if run.text == original_text:
            pass  # no match in this run — handled by cross-run pass if needed

    # ── Step 2: cross-run pass for PII that spans run boundaries ─────────────
    full_text = "".join(r.text for r in para.runs)
    needs_merge = any(orig in full_text for orig in replacement_map)

    if needs_merge:
        merged = full_text
        for original, replacement in replacement_map.items():
            merged = merged.replace(original, replacement)

        if merged != full_text and para.runs:
            # Write the fully-replaced text into the first run,
            # clear subsequent runs (preserves the first run's formatting).
            para.runs[0].text = merged
            for run in para.runs[1:]:
                run.text = ""


def process_docx(
    input_path: str,
    output_path: str,
    mode: str = "redact",
) -> dict[str, Any]:
    """
    Detect and sanitise PII from a DOCX file.

    Returns a summary dict.
    """
    doc: Document = Document(input_path)
    runs = _collect_runs(doc)
    paragraphs = _collect_paragraphs(doc)

    # ── 1. Build full text for pipeline ──────────────────────────────────────
    full_text = " ".join(r.text for r in runs if r.text)

    # ── 2. Run detection pipeline ─────────────────────────────────────────────
    analysis = pii_analyzer.analyze(full_text)
    enriched = context_analyzer.analyze(
        full_text,
        analysis["presidio_results"],
        analysis["indic_results"],
        analysis["label_pairs"],
    )
    deduped = confidence_scorer.deduplicate(enriched)
    scored = confidence_scorer.score_and_filter(deduped)
    to_mask = scored["to_mask"]

    # Ask masker for the replacement strings (use mode but also need raw map)
    mask_out = pii_masker.mask(full_text, to_mask, mode)

    # Build {original_value: replacement} map from results
    replacement_map: dict[str, str] = {}
    for result in to_mask:
        value = result.get("value", "")
        if not value:
            continue
        # Get the replacement for this specific span
        single_out = pii_masker.mask(value, [{**result, "start": 0, "end": len(value)}], mode)
        replacement_map[value] = single_out["masked_text"]

    # ── 3. Apply replacements paragraph-by-paragraph ──────────────────────────
    # Using paragraph-level replacement handles PII that spans multiple runs.
    for para in paragraphs:
        _replace_in_paragraph(para, replacement_map)

    # ── 4. Clear metadata PII ─────────────────────────────────────────────────
    props = doc.core_properties
    props.author = "Sanitized"
    props.last_modified_by = "Sanitized"

    doc.save(output_path)

    return {
        "pii_summary": confidence_scorer.get_summary(to_mask),
        "layer_breakdown": confidence_scorer.get_layer_breakdown(to_mask),
        "confidence_breakdown": {
            "high": scored["high_count"],
            "medium": scored["medium_count"],
        },
        "total_pii": len(to_mask),
    }

