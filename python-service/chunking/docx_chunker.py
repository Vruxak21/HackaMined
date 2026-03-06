"""
DOCX file chunker and merger using python-docx.

Splits a DOCX at element boundaries (paragraphs / tables) so the PII
detection pass can work on a reasonably-sized in-memory document.

Complexity notes
----------------
* Styles live in the document's ``/word/styles.xml`` part.  Rather than
  trying to clone the styles XML across documents (fragile), each chunk is
  created by **copying the entire original document** and then replacing
  every body element with only the elements belonging to that chunk.  This
  gives every chunk the complete style sheet, theme, fonts, and section
  properties of the original — no formatting is ever lost.

* Overlap elements are wrapped between plain-text marker paragraphs so the
  merger can strip them without any XML introspection.

Overlap markers (plain text paragraphs inserted by split, removed by merge)
---------------------------------------------------------------------------
    ##OVERLAP_BEFORE##
    <overlap elements from previous chunk>
    ##OVERLAP_END##
    ... main chunk content ...
    ##OVERLAP_AFTER##
    <overlap elements from next chunk>
    ##OVERLAP_END_AFTER##

Merge strategy
--------------
The original file is opened as the base document (preserving headers,
footers, styles, core properties, custom XML parts).  All body elements
except the final ``<w:sectPr>`` are cleared, then the sanitised content
from each chunk is inserted in order, skipping everything inside overlap
markers.
"""

from __future__ import annotations

import copy
import logging
import math
import os
import uuid
import tempfile
from pathlib import Path
from typing import List, Optional

from docx import Document  # type: ignore[import]
from docx.oxml.ns import qn  # type: ignore[import]
from lxml import etree  # type: ignore[import]

from chunking.chunk_models import ChunkMetadata, ChunkResult
from chunking.config import get_chunk_config

logger = logging.getLogger(__name__)

DEFAULT_CHUNK_SIZE = 400  # paragraphs per chunk

# Plain-text overlap markers (inserted as paragraph text)
_MARK_BEFORE_START = "##OVERLAP_BEFORE##"
_MARK_BEFORE_END = "##OVERLAP_END##"
_MARK_AFTER_START = "##OVERLAP_AFTER##"
_MARK_AFTER_END = "##OVERLAP_END_AFTER##"

_ALL_MARKERS = {
    _MARK_BEFORE_START,
    _MARK_BEFORE_END,
    _MARK_AFTER_START,
    _MARK_AFTER_END,
}


def _get_file_id(file_path: str) -> str:
    stem = Path(file_path).stem.replace("-", "").replace("_", "")
    if len(stem) >= 32 and all(c in "0123456789abcdefABCDEF" for c in stem[:32]):
        return stem[:32]
    return uuid.uuid4().hex


def _para_text(el: etree._Element) -> str:
    """Return the plain concatenated text of a <w:p> element."""
    return "".join(
        n.text or "" for n in el.iter(qn("w:t"))
    )


def _add_marker_para(body: etree._Element, text: str) -> None:
    """Insert a bare marker paragraph immediately before <w:sectPr>."""
    # Create the element without a parent first, then insert in the correct
    # position.  Using etree.SubElement would append *after* <w:sectPr> which
    # produces invalid OOXML and causes merge to treat overlap content as main.
    p = etree.Element(qn("w:p"))
    r = etree.SubElement(p, qn("w:r"))
    t = etree.SubElement(r, qn("w:t"))
    t.text = text
    _insert_before_sectpr(body, p)


def _body_content_elements(doc: Document) -> List[etree._Element]:
    """
    Return all direct body children except <w:sectPr> (section properties).
    Includes <w:p> and <w:tbl> in document order.
    """
    body = doc.element.body
    return [
        child for child in body
        if child.tag not in (qn("w:sectPr"),)
    ]


def _clear_body_content(doc: Document) -> None:
    """Remove all body children except <w:sectPr>."""
    body = doc.element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def _insert_before_sectpr(body: etree._Element, el: etree._Element) -> None:
    """Insert *el* into *body* immediately before <w:sectPr> (or at end)."""
    sect_pr = body.find(qn("w:sectPr"))
    if sect_pr is not None:
        sect_pr.addprevious(el)
    else:
        body.append(el)


class DOCXChunker:
    """Splits large DOCX files into element-boundary-aligned chunks."""

    # ------------------------------------------------------------------
    # split
    # ------------------------------------------------------------------

    def split(self, file_path: str, chunk_size: int | None = None) -> List[ChunkMetadata]:
        """
        Open *file_path*, collect all body elements (paragraphs + tables)
        in document order, group them into chunks, and write one temp DOCX
        per chunk.

        Parameters
        ----------
        chunk_size:
            Paragraphs per chunk. Falls back to DEFAULT_CHUNK_SIZE if None.
        """
        paras_per_chunk: int = chunk_size or DEFAULT_CHUNK_SIZE
        overlap: int = 2  # overlap paragraphs

        doc = Document(file_path)
        all_elements = _body_content_elements(doc)
        total_elements = len(all_elements)

        if total_elements == 0:
            logger.warning("DOCXChunker.split: no body elements found in %s", file_path)
            return []

        total_chunks = math.ceil(total_elements / paras_per_chunk)
        file_id = _get_file_id(file_path)
        tmp_dir = Path(tempfile.gettempdir())
        chunk_list: List[ChunkMetadata] = []

        for chunk_idx in range(total_chunks):
            start_el = chunk_idx * paras_per_chunk
            end_el = min(start_el + paras_per_chunk, total_elements)

            temp_input_path = str(tmp_dir / f"chunk_{file_id}_{chunk_idx}.docx")
            temp_output_path = str(tmp_dir / f"chunk_{file_id}_{chunk_idx}_out.docx")

            # ── Build chunk document ──────────────────────────────────────
            # Open a fresh copy of the original so styles/themes/fonts are
            # identical — avoids cross-document style cloning entirely.
            chunk_doc = Document(file_path)
            _clear_body_content(chunk_doc)
            chunk_body = chunk_doc.element.body

            # ── Overlap BEFORE ────────────────────────────────────────────
            if chunk_idx > 0 and overlap > 0:
                ov_start = max(0, start_el - overlap)
                _add_marker_para(chunk_body, _MARK_BEFORE_START)
                for el in all_elements[ov_start:start_el]:
                    _insert_before_sectpr(chunk_body, copy.deepcopy(el))
                _add_marker_para(chunk_body, _MARK_BEFORE_END)

            # ── Main content ──────────────────────────────────────────────
            for el in all_elements[start_el:end_el]:
                _insert_before_sectpr(chunk_body, copy.deepcopy(el))

            # ── Overlap AFTER ─────────────────────────────────────────────
            if chunk_idx < total_chunks - 1 and overlap > 0:
                ov_end = min(end_el + overlap, total_elements)
                _add_marker_para(chunk_body, _MARK_AFTER_START)
                for el in all_elements[end_el:ov_end]:
                    _insert_before_sectpr(chunk_body, copy.deepcopy(el))
                _add_marker_para(chunk_body, _MARK_AFTER_END)

            chunk_doc.save(temp_input_path)

            chunk_list.append(
                ChunkMetadata(
                    chunk_index=chunk_idx,
                    total_chunks=total_chunks,
                    file_type="docx",
                    start_boundary=start_el,
                    end_boundary=end_el,
                    overlap_before=overlap if chunk_idx > 0 else 0,
                    overlap_after=overlap if chunk_idx < total_chunks - 1 else 0,
                    temp_input_path=temp_input_path,
                    temp_output_path=temp_output_path,
                )
            )

        logger.info(
            "DOCXChunker.split: %d elements → %d chunks  (file=%s)",
            total_elements,
            total_chunks,
            file_path,
        )
        return chunk_list

    # ------------------------------------------------------------------
    # merge
    # ------------------------------------------------------------------

    def merge(
        self,
        chunk_results: List[ChunkResult],
        chunk_output_paths: List[str],
        output_path: str,
        original_file_path: str,
    ) -> None:
        """
        Reassemble processed chunk DOCX files into a single document.

        Uses the original file as the base (preserving headers, footers,
        styles, core properties) and replaces its body content with the
        sanitised content from each chunk, in order, skipping all overlap
        marker regions.
        """
        ordered = sorted(
            zip(chunk_results, chunk_output_paths),
            key=lambda pair: pair[0].chunk_index,
        )

        output_doc = Document(original_file_path)
        _clear_body_content(output_doc)
        out_body = output_doc.element.body

        para_count = 0
        table_count = 0

        for result, chunk_path in ordered:
            try:
                chunk_doc = Document(chunk_path)
            except Exception as exc:
                logger.warning(
                    "DOCXChunker.merge: cannot open chunk %d (%s): %s",
                    result.chunk_index,
                    chunk_path,
                    exc,
                )
                continue

            skip = False  # True while inside an overlap region

            for el in _body_content_elements(chunk_doc):
                tag_local = el.tag.split("}")[-1] if "}" in el.tag else el.tag

                # Check if this element is a marker paragraph
                if tag_local == "p":
                    text = _para_text(el).strip()
                    if text in (_MARK_BEFORE_START, _MARK_AFTER_START):
                        skip = True
                        continue
                    if text in (_MARK_BEFORE_END, _MARK_AFTER_END):
                        skip = False
                        continue

                if skip:
                    continue

                # Real content — deep-copy into output document
                _insert_before_sectpr(out_body, copy.deepcopy(el))

                if tag_local == "p":
                    para_count += 1
                elif tag_local == "tbl":
                    table_count += 1

        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        output_doc.save(output_path)

        # ── Integrity verification ────────────────────────────────────────
        try:
            orig_doc = Document(original_file_path)
            orig_elements = _body_content_elements(orig_doc)
            orig_para_count = sum(
                1 for e in orig_elements
                if (e.tag.split("}")[-1] if "}" in e.tag else e.tag) == "p"
            )
            orig_table_count = sum(
                1 for e in orig_elements
                if (e.tag.split("}")[-1] if "}" in e.tag else e.tag) == "tbl"
            )
            if orig_para_count != para_count or orig_table_count != table_count:
                logger.warning(
                    "DOCXChunker.merge: element count mismatch — "
                    "original(paras=%d, tables=%d) output(paras=%d, tables=%d)  (file=%s)",
                    orig_para_count,
                    orig_table_count,
                    para_count,
                    table_count,
                    output_path,
                )
            else:
                logger.info(
                    "DOCX merged: %d paragraphs, %d tables preserved  (file=%s)",
                    para_count,
                    table_count,
                    output_path,
                )
        except Exception:
            pass

    # ------------------------------------------------------------------
    # copy_styles  (kept as documented helper; strategy uses full-doc copy)
    # ------------------------------------------------------------------

    @staticmethod
    def copy_styles(source_doc: Document, target_doc: Document) -> None:
        """
        Copy all styles from *source_doc* into *target_doc*.

        In practice, :meth:`split` avoids needing this by cloning the
        entire source document for each chunk; this helper is provided
        for any caller that builds a fresh ``Document()`` and needs to
        import styles manually.
        """
        src_styles_el = source_doc.element.find(
            ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}styles"
        )
        tgt_styles_el = target_doc.element.find(
            ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}styles"
        )
        if src_styles_el is None or tgt_styles_el is None:
            return

        # Replace all style children in target with copies from source
        for child in list(tgt_styles_el):
            tgt_styles_el.remove(child)
        for child in src_styles_el:
            tgt_styles_el.append(copy.deepcopy(child))

    # ------------------------------------------------------------------
    # cleanup
    # ------------------------------------------------------------------

    def cleanup(self, chunk_metadata_list: List[ChunkMetadata]) -> None:
        """Delete all temporary files created by :meth:`split`."""
        for meta in chunk_metadata_list:
            for path in (meta.temp_input_path, meta.temp_output_path):
                try:
                    os.remove(path)
                except OSError:
                    pass


# ── Module-level chunked processing function ──────────────────────────────────

def _noop_progress(chunk_idx: int, status: str) -> None:  # noqa: ARG001
    pass


def _collect_all_paragraphs_with_tables(doc) -> list:
    """
    Return every paragraph from body and table cells.
    Each item is a python-docx Paragraph object.
    """
    paras = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paras.extend(cell.paragraphs)
    for section in doc.sections:
        for container in (section.header, section.footer):
            if not container.is_linked_to_previous:
                paras.extend(container.paragraphs)
    return paras


def _replace_para_text(para, replacement_map: dict[str, str]) -> None:
    """Apply replacement_map to a paragraph, handling cross-run PII."""
    if not para.runs:
        return
    # Per-run pass (preserves formatting for within-run PII)
    for run in para.runs:
        for original, replacement in replacement_map.items():
            if original in run.text:
                run.text = run.text.replace(original, replacement)
    # Cross-run pass (handles PII spanning run boundaries)
    full = "".join(r.text for r in para.runs)
    modified = full
    for original in sorted(replacement_map, key=len, reverse=True):
        modified = modified.replace(original, replacement_map[original])
    if modified != full and para.runs:
        para.runs[0].text = modified
        for run in para.runs[1:]:
            run.text = ""


def process_docx_chunked(
    input_path: str,
    output_path: str,
    mode: str = "redact",
    config: dict | None = None,
    progress_cb=_noop_progress,
) -> dict:
    """
    End-to-end chunked DOCX processing with the new pipeline.

    1. Open DOCX with python-docx, collect all paragraphs + table cells.
    2. Group paragraphs into chunks of config["chunk_size"].
    3. Run detect_pii_batch() on chunk texts in parallel.
    4. Apply masking SEQUENTIALLY (python-docx not thread-safe for writes).
    5. Save once at the end.
    """
    import threading
    from concurrent.futures import ThreadPoolExecutor, as_completed
    from pathlib import Path

    from detection.masker import pii_masker
    from pipeline.detector import detect_pii_batch

    if config is None:
        config = {"use_regex": True, "use_spacy": True, "use_bert": False,
                  "spacy_model": "en_core_web_sm", "chunk_size": DEFAULT_CHUNK_SIZE,
                  "workers": 4}

    paras_per_chunk = config.get("chunk_size", DEFAULT_CHUNK_SIZE)
    workers = config.get("workers", 4)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)

    doc = Document(input_path)
    all_paras = _collect_all_paragraphs_with_tables(doc)

    if not all_paras:
        doc.save(output_path)
        return _build_docx_summary([], 0)

    # ── Group paragraphs into chunks ─────────────────────────────────
    para_groups: list[list] = [
        all_paras[i : i + paras_per_chunk]
        for i in range(0, max(1, len(all_paras)), paras_per_chunk)
    ]
    total = len(para_groups)

    for i in range(total):
        progress_cb(i, "pending")

    # ── Extract text per chunk for detection ──────────────────────────
    chunk_texts: list[str] = []
    for group in para_groups:
        chunk_texts.append(
            "\n".join(" ".join(r.text for r in p.runs) for p in group)
        )

    # ── Detection: parallel across chunks ─────────────────────────────
    lock = threading.Lock()
    det_results: dict[int, tuple[dict[str, str], list[dict]]] = {}
    all_detections: list[dict] = []

    def _detect_group(idx: int, text: str):
        progress_cb(idx, "processing")
        try:
            dets_list = detect_pii_batch([text], config)
            dets = dets_list[0]

            repl_map: dict[str, str] = {}
            for det in dets:
                val = det["value"]
                if not val or val in repl_map:
                    continue
                if mode == "redact":
                    repl_map[val] = "[REDACTED]"
                elif mode == "mask":
                    repl_map[val] = pii_masker.get_partial_mask(
                        val, det["entity_type"]
                    )
                elif mode == "tokenize":
                    repl_map[val] = pii_masker.get_token(det["entity_type"])
                else:
                    repl_map[val] = "[REDACTED]"

            progress_cb(idx, "done")
            return idx, repl_map, dets
        except Exception:
            progress_cb(idx, "failed")
            raise

    with ThreadPoolExecutor(max_workers=min(workers, total)) as executor:
        futures = {
            executor.submit(_detect_group, i, chunk_texts[i]): i
            for i in range(total)
        }
        for future in as_completed(futures):
            idx, repl_map, dets = future.result()
            with lock:
                det_results[idx] = (repl_map, dets)
                all_detections.extend(dets)

    # ── Apply masking SEQUENTIALLY (python-docx writes not thread-safe) ──
    for i, group in enumerate(para_groups):
        repl_map, _ = det_results.get(i, ({}, []))
        if repl_map:
            for para in group:
                _replace_para_text(para, repl_map)

    # Clean metadata
    core = doc.core_properties
    for attr in ("author", "last_modified_by", "comments"):
        try:
            setattr(core, attr, "Sanitized")
        except Exception:
            pass

    doc.save(output_path)
    return _build_docx_summary(all_detections, total)


def _build_docx_summary(detections: list[dict], chunk_count: int) -> dict:
    pii_summary: dict[str, int] = {}
    layer_breakdown: dict[str, int] = {"regex": 0, "spacy": 0, "bert": 0}
    high = medium = 0

    for det in detections:
        et = det["entity_type"]
        pii_summary[et] = pii_summary.get(et, 0) + 1
        layer = det.get("layer", "regex")
        if layer in layer_breakdown:
            layer_breakdown[layer] += 1
        else:
            layer_breakdown[layer] = 1
        score = det.get("score", 0)
        if score >= 0.85:
            high += 1
        elif score >= 0.60:
            medium += 1

    return {
        "pii_summary":          pii_summary,
        "total_pii":            len(detections),
        "layer_breakdown":      layer_breakdown,
        "confidence_breakdown": {"high": high, "medium": medium},
        "chunk_count":          chunk_count,
    }
